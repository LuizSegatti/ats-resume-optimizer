# Updated function "Apply Replacements to DOCX"
# Resume Matcher App – GPT-Based Analysis, Resume Rewriter, and Excel Logger (Streamlit Version)

import streamlit as st
import fitz  # PyMuPDF
import docx
import os
import re
import openpyxl
import json
import tempfile
import time
from datetime import datetime
from gpt_helper_work_version import get_resume_analysis, generate_cover_letter
from dotenv import load_dotenv
from openpyxl.utils import get_column_letter
from openpyxl.worksheet.table import Table

# Load API Key from .env
load_dotenv()
api_key = os.getenv("OPENAI_API_KEY")

# === Save Customized Cover Letter ===
def save_customized_cover_letter(template_path, output_folder, cover_text, resume_text, company_name):
    from docx import Document

    candidate_info = {}

    # Extract candidate data using regex
    name_match = re.search(r"^\s*([A-Z][a-z]+(?:\s+[A-Z][a-z]+)+)", resume_text, re.MULTILINE)
    email_match = re.search(r"[\w\.-]+@[\w\.-]+", resume_text)
    phone_match = re.search(r"(\(?\d{3}\)?[-.\s]?\d{3}[-.\s]?\d{4})", resume_text)
    linkedin_match = re.search(r"(https?://(www\.)?linkedin\.com/in/[\w\d-]+)", resume_text)
    city_state_match = re.search(r"\b([A-Z][a-z]+(?:\s[A-Z][a-z]+)*,\s*[A-Z]{2})\b", resume_text)

    if name_match:
        candidate_info["[Candidate Name]"] = name_match.group(1).strip()
    if email_match:
        candidate_info["[candidate email]"] = email_match.group(0).strip()
    if phone_match:
        digits = re.sub(r"\D", "", phone_match.group(0))
        if len(digits) == 10:
            candidate_info["[candidate phone]"] = f"({digits[:3]}) {digits[3:6]}-{digits[6:]}"
    if linkedin_match:
        candidate_info["[LinkedIn]"] = linkedin_match.group(1).strip()
    if city_state_match:
        candidate_info["[candidate city, State]"] = city_state_match.group(1).strip()

    candidate_info["[date]"] = datetime.today().strftime("%B %d, %Y")

    doc = Document(template_path)

    for para in doc.paragraphs:
        for run in para.runs:
            for placeholder, value in candidate_info.items():
                if placeholder in run.text:
                    run.text = run.text.replace(placeholder, value)

    candidate_name = candidate_info.get("[Candidate Name]", "Candidate")
    filename = f"CV_{candidate_name}-{company_name}-{datetime.today().strftime('%Y-%m-%d')}.docx"
    filepath = os.path.join(output_folder, filename)
    doc.save(filepath)
    return filepath, candidate_name

# === Extract Resume Text ===
def extract_text(file_path):
    if file_path.lower().endswith(".pdf"):
        doc = fitz.open(file_path)
        text = "\n".join([page.get_text() for page in doc])
        doc.close()
        return text
    elif file_path.lower().endswith(".docx"):
        while True:
            try:
                return "\n".join([p.text for p in docx.Document(file_path).paragraphs])
            except Exception:
                raise Exception(f"Please make sure the file is not open: {file_path}")
    return ""

# === Parse Replacements from GPT Output ===
def parse_replacements(gpt_output):
    json_block = re.search(r"```json\s*(\[.*?\])\s*```", gpt_output, re.DOTALL)
    if json_block:
        try:
            replacements = json.loads(json_block.group(1))
            return [(item["was"], item["new"], item.get("section", "Unknown")) for item in replacements]
        except Exception as e:
            print("⚠️ Failed to parse JSON change log:", e)
            return []
    return re.findall(r"(?i)replace [“\"](.+?)[”\"] with [“\"](.+?)[”\"]", gpt_output)

# === Apply Replacements to DOCX ===
def apply_replacements_to_docx(original_path, replacements, save_path=None):
    while True:
        try:
            doc = docx.Document(original_path)
            break
        except Exception:
            st.warning(f"⚠️ Please close the resume file:\n{original_path}")
            time.sleep(1)
    changes = []
    for para in doc.paragraphs:
        for old, new in replacements:
            if old.lower() in para.text.lower():
                pattern = re.compile(re.escape(old), re.IGNORECASE)
                updated_text = pattern.sub(new, para.text)
                if updated_text != para.text:
                    para.text = updated_text
                    changes.append((old, new, "Resume Body"))

    # Save if save_path provided
    if save_path:
        doc.save(save_path)
    return doc, save_path or original_path

# === Extract Final Optimized Resume Text ===
def extract_final_resume_text(gpt_output):
    block = re.search(r"(?:## Final optimized resume|```text)(.*?)(?:```|\Z)", gpt_output, re.DOTALL | re.IGNORECASE)
    if block:
        return block.group(1).strip()
    return None

# === Extract Company Name from GPT Output ===
def extract_company_name_from_gpt(gpt_output):
    match = re.search(r"(?i)Company(?: Name)?:\s*(.+)", gpt_output)
    if match:
        return match.group(1).strip()
    return "UnknownCompany"

# === Log Results into Excel Tracker ===
def log_gpt_results(tracker_path, resume_name, jd_name, score, changes, resume_filename, company_name):
    wb = None
    while wb is None:
        try:
            wb = openpyxl.load_workbook(tracker_path)
        except PermissionError:
            st.warning("⚠️ Please close the Excel file before continuing.")
            time.sleep(1)

    try:
        today = datetime.today()
        job_title = jd_name.replace(".docx", "")
        resume_version = resume_filename.replace(".docx", "")

        def format_score(score_text):
            try:
                return float(score_text.replace('%', '').strip()) / 100
            except:
                return None

        # ATS_Report_Log
        ats = wb["ATS_Report_Log"]
        ats_table = ats.tables["ATS_Report"]
        ats.append([
            ats.max_row,
            resume_version,
            job_title,
            company_name,
            format_score(score),
            f"{len(changes)} changes",
            today.date()
        ])
        ats_table.ref = f"A1:{get_column_letter(ats.max_column)}{ats.max_row}"

        # Change_Log
        chg = wb["Change_Log"]
        chg_table = chg.tables["Change_Log"]
        for old, new, section in changes:
            chg.append([
                chg.max_row,
                resume_version,
                old,
                new,
                section,
                job_title,
                today.date()
            ])
        chg_table.ref = f"A1:{get_column_letter(chg.max_column)}{chg.max_row}"

        # Resume_Inventory
        inv = wb["Resume_Inventory"]
        inv_table = inv.tables["Resume_Inventory"]
        inv.append([
            resume_version,
            resume_name,
            job_title,
            f"/03-Customized_Resumes/{resume_filename}",
            today.date()
        ])
        inv_table.ref = f"A1:{get_column_letter(inv.max_column)}{inv.max_row}"

        # Job_Application_Tracker (if exists)
        if "Job_Application_Tracker" in wb.sheetnames:
            tracker = wb["Job_Application_Tracker"]
            tracker_table = tracker.tables["tblJobApplications"]
            tracker.append([
                tracker.max_row,
                resume_version,
                job_title,
                company_name,
                today.date(),
                "Pending",
                "Auto-logged"
            ])
            tracker_table.ref = f"A1:{get_column_letter(tracker.max_column)}{tracker.max_row}"

        wb.save(tracker_path)

    finally:
        wb.close()

